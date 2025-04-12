from docx import Document

class DocxTranslator:
    def __init__(self, translation_manager):
        self.translation_manager = translation_manager

    def translate_docx(self, input_path, lang_pair, progress_callback=None):
        # Charger le document original
        doc = Document(input_path)
        new_doc = Document()
        
        total_elements = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
        current_element = 0
        
        # Traduire les paragraphes
        for para in doc.paragraphs:
            if para.text.strip():
                translated_text = self.translation_manager.translate(para.text.strip(), lang_pair)
                new_doc.add_paragraph(translated_text)
            else:
                new_doc.add_paragraph()
            
            current_element += 1
            if progress_callback:
                progress_callback(int(current_element * 100 / total_elements))
        
        # Traduire les tableaux
        for table in doc.tables:
            # Créer un nouveau tableau avec le même nombre de colonnes
            new_table = new_doc.add_table(rows=0, cols=len(table.columns))
            
            for row in table.rows:
                # Ajouter une nouvelle ligne
                new_cells = new_table.add_row().cells
                
                # Traduire chaque cellule
                for i, cell in enumerate(row.cells):
                    if cell.text.strip():
                        translated_cell = self.translation_manager.translate(cell.text.strip(), lang_pair)
                        new_cells[i].text = translated_cell
                
                current_element += 1
                if progress_callback:
                    progress_callback(int(current_element * 100 / total_elements))
        
        # Sauvegarder le document traduit
        output_path = input_path.rsplit('.', 1)[0] + '_translated.docx'
        new_doc.save(output_path)
        return output_path